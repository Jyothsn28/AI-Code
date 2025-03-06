
from fastapi import logger
import fitz
from langchain_community.document_loaders import PyPDFLoader
from sklearn.metrics.pairwise import cosine_similarity
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.callbacks.manager import get_openai_callback
from langchain.embeddings import AzureOpenAIEmbeddings
from bson.objectid import ObjectId
from pymongo.mongo_client import MongoClient
# from text_table_extract import get_text_tables
import logging
# import fitz
from collections import defaultdict
import datetime
# import spacy
import hashlib
import math

from docx import Document
from pymongo.mongo_client import MongoClient

from langchain_community.chat_models import AzureChatOpenAI
mongo_uri = "mongodb+srv://chaitanya:BVjn9TRZ4gu4Bs4X@cluster0.80157cr.mongodb.net/?retryWrites=true&w=majority"

upload_files_limit = 2

dbobj = None
docs_coll  = None
cksum_coll = None
uri = mongo_uri
# global dbobj
# global docs_coll
# global cksum_coll
client = MongoClient(uri)
    # Send a ping to confirm a successful connection
try:
    client.admin.command('ping')
    print("Pinged your deployment. You successfully connected to MongoDB!")
except Exception as e:
        print(e)



dbobj = client['ams_service']
docs_coll  = dbobj['ams_service_doc']
cksum_coll = dbobj['ams_service_cksum']
def doc_loaded(dbobj, cksum):
    found = False
    for x in cksum_coll.find({ "cksum": cksum}):
        print(x)
        found = True
    return found

def extract_text_pdf(downloaded, document_name, location, checksum, chunk=30, overlap=20):
    """Module to extract and chunk texts from a given PDF
    Parameters:
        - documents: A list of paths to the PDF files whose text needs to be extracted
        - chunk: Maximum number of sentences to store in 1 single chunk
    Returns:
        The module returns a dictionary containing all the text and page numbers from the given PDFs
    """
    all_text = defaultdict(list)
    #texts = []
    #page_nums = []
    # document = fitz.open('pdf', downloaded)
    document = downloaded
    if not document:
        return all_text
 
    create_date = get_doc_creation_date(document)
    logger.debug(f"document create date ={create_date}, {type(create_date)}")
 
    if not document_name:
        document_name = "no_name_" + str(uuid.uuid4())
 
    first_half = list()
    second_half = list()
    split_sentences = list()
 
    text_data = list()

    text_data,tables_data = text_from_fitz_output(document)
   
    for td in text_data:
        page_no = td['page_no']
        sentences = td['sentences']

        if len(sentences) > 8:
            first_half = sentences[: len(sentences) // 2]
        else:
            first_half = sentences[:]

        logger.info(f"Analysing page {page_no}")
        logger.debug(f"Text in this page = \n{sentences}")
        if len(second_half):
            split_sentences = second_half + first_half
            logger.debug(f"first half this page = \n{first_half}")
            logger.debug(f"second half previous page = \n{second_half}")
            logger.debug(f"split_sentences = \n{split_sentences}")
 
        logger.debug(f"len of sentences = {len(sentences)}")
        logger.debug(f"len of split sentences = {len(split_sentences)}")
        for i in range(0, len(sentences), overlap):
            logger.debug(f"sentence part = {i}")
            logger.debug(". ".join(sentences[i : i + chunk]))

        if len(second_half):
            for i in range(0, len(split_sentences), overlap):
                logger.debug(f"split sentence part = {i}")
                logger.debug(". ".join(split_sentences[i : i + chunk]))
        logger.debug("-------")
 
        if len(second_half):
            all_text[document_name].extend(
                [
                    (
                        {
                            #"Content": ". ".join(split_sentences[i : i + chunk]),
                            "Content": split_sentences[i : i + chunk],
                            "metadata": {
                                "source": document_name,
                                "creation_date": create_date,
                                "page_no": page_no,
                                "paragraph_no": i,
                                "location": location,
                                "checksum": checksum
                            },
                        }
                    )
                    for i in range(0, len(split_sentences), overlap)
                ]
            )
        else:
            all_text[document_name].extend(
                [
                    (
                        {
                            # "Content": ". ".join(sentences[i : i + chunk]),
                            "Content": sentences[i: i + chunk],
                            "metadata": {
                                "source": document_name,
                                "creation_date": create_date,
                                "page_no": page_no,
                                "paragraph_no": i,
                                "location": location,
                                "checksum": checksum
                            },
                        }
                    )
                    for i in range(0, len(sentences), overlap)
                ]
            )

        if len(sentences) > 8:
            second_half = sentences[len(sentences) // 2 :]
        else:
            if len(sentences) > 5:
                second_half = sentences[-3:]
            elif len(sentences) > 1:
                second_half = sentences[-1 * math.floor(len(sentences) / 2):]
            else:
                second_half = list()

        logger.debug(f"page = {page_no} done")
        #print(f"first half = {first_half}")
        #print(f"second half = {second_half}")
        logger.debug("------------------------------------------")
        # break

    for tab_data in tables_data:
        all_text[document_name].append(
                (
                    {
                        # "Content": ". ".join(split_sentences[i : i + chunk]),
                        "Content": [tab_data['table']],
                        "metadata": {
                            "source": document_name,
                            "creation_date": create_date,
                            "page_no": tab_data['page_no'],
                            "paragraph_no": '',
                            "location": location,
                            "checksum": checksum
                        },
                    }
                )
        )

    extracted_text = dict(all_text)
    return extracted_text

def extract_text(dbobj, uploaded_docs_list, filename):
    extracted_texts = list()
    checksums = list()
    for uploaded_doc in uploaded_docs_list:
        # document = fitz.open("pdf", uploaded_doc.getvalue())
        # file_data = uploaded_doc.getvalue()
        # file_name = uploaded_doc.name

        file_data = fitz.open(stream = uploaded_doc)
        chksum_text = ''.join([file_data[pg_num].get_text() for pg_num in range(0,file_data.page_count)]).encode("utf-8")
        file_name = filename


        print(f"reading doc {file_name}")
        cksum = get_checksum(chksum_text)
        if doc_loaded(dbobj, cksum):
            print(f"{file_name} already loaded {cksum}, skipping")
            continue
        print(f"File {file_name} {cksum} not already loaded")
        extracted = extract_text_pdf(file_data, file_name, "", cksum)
        extracted_texts.append(extracted)

        print("checksum value =", {'name': file_name, 'cksum': cksum})
        checksums.append({'name': file_name, 'cksum': cksum})
    return [extracted_texts, checksums]

# def extract_text_from_docx(dbobj, uploaded_docs_list, filename):
#     # doc = Document(uploaded_docs_list)
#     doc = uploaded_docs_list
#     print("doc load", doc)
#     # doc_text = read_docx(doc)
#     doc_text = doc
#     text = []
#     file_name = filename
#     file_already_uploaded = False
#          # Access paragraphs
#     for para in doc.paragraphs:
#         print(para.text) 
            
#         text.append(para.text)
#     # for paragraph in doc_text.paragraphs:
#     #     text.append(paragraph.text)
#     print("text",text)
#     print(f"reading doc {file_name}")
#     chksum_text = ''.join([text[pg_num].get_text() for pg_num in range(0,text.page_count)]).encode("utf-8")
  
#     cksum = get_checksum(chksum_text)
    
    
#     return    [text, cksum, file_already_uploaded]
#     # return "\n".join(text)

    
# def extract_text(dbobj, uploaded_docs_list, filename):
#     extracted_texts = list()
#     checksums = list()
#     for uploaded_doc in uploaded_docs_list:
        

#         file_data = fitz.open(stream = uploaded_doc)
#         chksum_text = ''.join([file_data[pg_num].get_text() for pg_num in range(0,file_data.page_count)]).encode("utf-8")
#         file_name = filename
#         file_already_uploaded = False
#         print(f"reading doc {file_name}")
#         cksum = get_checksum(chksum_text)
#         if doc_loaded(dbobj, cksum):
#             file_already_uploaded = True
#             print(f"{file_name} already loaded {cksum}, skipping")
#             continue
#         print(f"File {file_name} {cksum} not already loaded")
        
#         extracted = extract_text_pdf(file_data, file_name, "", cksum)
#         extracted_texts.append(extracted)

#         print("checksum value =", {'name': file_name, 'cksum': cksum})
#         # checksums.append({'name': file_name, 'cksum': cksum, 'doc_uri': gsutil_uri})
        

#     return [extracted_texts, cksum, file_already_uploaded]

def get_checksum(data):
    cksum = hashlib.md5(data).hexdigest()
    return cksum
def load_to_db(dbobj, docs, checksums):
    for doc in docs:
        for k, v in doc.items():
            #print(f"-----------------------\nInserting {v}\n----------------------------")
            docs_coll.insert_many(v)
embeddings = AzureOpenAIEmbeddings(
    deployment="embed_small",
    azure_endpoint="https://genai-cube.openai.azure.com/",
    openai_api_type="azure",
    openai_api_key = "0e1c91cf40fc4432a9da9d3d4b493dc4",
    chunk_size= 2560
)
def embed_doc(docs_extracted_texts):
    final_embed = list()
    for doc_elem in docs_extracted_texts:
        embedded_docs = dict()
        for doc, extracted in doc_elem.items():
            embedded_extracts = list()
            for split in extracted:
                embedded_split = dict()
                page_content = " ".join(split['Content'])
                embedded_page_content = embeddings.embed_query(page_content)
                embedded_split['Content'] = page_content
                embedded_split['Content_vector'] = embedded_page_content
                embedded_split['metadata'] = split['metadata']
                embedded_extracts.append(embedded_split)
            embedded_docs[doc] = embedded_extracts
        final_embed.append(embedded_docs)

    return final_embed
def write_to_db(uploaded_files, filename):
   

    [all_docs_texts, checksums, file_already_upload] = extract_text_from_docx(dbobj, uploaded_files, filename)
    print(len(all_docs_texts))
    #print(all_docs_texts)
    print("embedding text")
    embedded_docs = embed_doc(all_docs_texts)
    #print(embedded_docs)
    print("writing to db")
    load_to_db(dbobj, embedded_docs, checksums)
    return [file_already_upload, checksums]